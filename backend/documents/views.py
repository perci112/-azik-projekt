from django.shortcuts import render
from rest_framework import generics, status, permissions
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, AllowAny
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.utils import timezone
from django.http import JsonResponse, FileResponse
from django.core.files.base import ContentFile
import io
import zipfile
from datetime import datetime
from django.db import transaction
from django.views.decorators.csrf import ensure_csrf_cookie
import json
from docx import Document as DocxDocument
import mammoth  # konwersja .docx -> HTML

from .models import (
    UserProfile, Document, EditableField, 
    DocumentAssignment, FieldValue, DocumentVersion
)
from .serializers import (
    UserSerializer, DocumentSerializer, EditableFieldSerializer,
    DocumentAssignmentSerializer, FieldValueSerializer, DocumentVersionSerializer,
    LoginSerializer, DocumentUploadSerializer, FieldCreationSerializer,
    AssignDocumentSerializer, SubmitFieldValuesSerializer
)
@api_view(['GET'])
@permission_classes([AllowAny])
@ensure_csrf_cookie
def csrf_token(request):
    """Ustawia ciasteczko CSRF dla aplikacji SPA"""
    return Response({"detail": "CSRF cookie set"})



@api_view(['POST'])
@permission_classes([AllowAny])
def login_view(request):
    """Endpoint do logowania użytkowników"""
    serializer = LoginSerializer(data=request.data)
    if serializer.is_valid():
        username = serializer.validated_data['username']
        password = serializer.validated_data['password']
        
        user = authenticate(request, username=username, password=password)
        if user and user.is_active:
            login(request, user)
            
            # Pobierz lub stwórz profil użytkownika
            profile, created = UserProfile.objects.get_or_create(
                user=user,
                defaults={'role': 'admin' if user.is_superuser else 'user'}
            )
            
            # Ensure correct role for superuser
            if user.is_superuser and profile.role != 'admin':
                profile.role = 'admin'
                profile.save()
            
            user_data = UserSerializer(user).data
            return Response({
                'success': True,
                'user': user_data,
                'message': 'Zalogowano pomyślnie'
            })
        else:
            return Response({
                'success': False,
                'message': 'Nieprawidłowe dane logowania'
            }, status=status.HTTP_401_UNAUTHORIZED)
    else:
        return Response({
            'success': False,
            'message': 'Błędne dane wejściowe',
            'errors': serializer.errors
        }, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def logout_view(request):
    """Endpoint do wylogowania"""
    logout(request)
    return Response({'success': True, 'message': 'Wylogowano pomyślnie'})


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def current_user(request):
    """Endpoint do pobrania danych aktualnego użytkownika"""
    serializer = UserSerializer(request.user)
    return Response(serializer.data)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def complete_profile(request):
    """Uzupełnij dane profilu przy pierwszym logowaniu (Discord lub inne)."""
    data = request.data or {}
    first_name = data.get('first_name', '').strip()
    last_name = data.get('last_name', '').strip()
    index = data.get('index', '').strip()
    section = data.get('section', '').strip()

    if not first_name or not last_name or not index or not section:
        return Response({'error': 'Wymagane: imię, nazwisko, index i sekcja.'}, status=status.HTTP_400_BAD_REQUEST)

    request.user.first_name = first_name
    request.user.last_name = last_name
    request.user.save()

    profile, _ = UserProfile.objects.get_or_create(user=request.user, defaults={'role': 'user'})
    profile.index = index
    profile.section = section
    profile.profile_completed = True
    profile.save()

    return Response({'success': True})


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def users_list(request):
    """Lista wszystkich użytkowników (tylko dla adminów)"""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)
    # Admin może zarządzać tylko swoją sekcją; musi mieć ustawioną sekcję
    admin_section = (user_profile.section or '').strip()
    if not admin_section:
        return Response({'error': 'Administrator nie ma ustawionej sekcji w profilu.'}, status=status.HTTP_400_BAD_REQUEST)

    from django.db.models import Q
    # Użytkownicy z tej samej sekcji ORAZ dodatkowo sam admin (możliwość przypisania do siebie)
    users_qs = User.objects.filter(
        Q(userprofile__role='user', userprofile__section=admin_section) | Q(id=request.user.id)
    ).distinct().order_by('username')
    serializer = UserSerializer(users_qs, many=True)
    return Response(serializer.data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def users_all(request):
    """Lista wszystkich użytkowników (tylko superuser)."""
    if not request.user.is_superuser:
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)
    users = User.objects.all().order_by('username')
    serializer = UserSerializer(users, many=True)
    return Response(serializer.data)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def set_user_role(request, user_id: int):
    """Nadanie roli admin/user (tylko superuser)."""
    if not request.user.is_superuser:
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)

    role = (request.data or {}).get('role')
    if role not in ('admin', 'user'):
        return Response({'error': 'Nieprawidłowa rola'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        target = User.objects.get(id=user_id)
    except User.DoesNotExist:
        return Response({'error': 'Użytkownik nie istnieje'}, status=status.HTTP_404_NOT_FOUND)

    profile, _ = UserProfile.objects.get_or_create(user=target, defaults={'role': 'user'})
    profile.role = role
    profile.save()
    return Response({'success': True})


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def upload_document(request):
    """Upload dokumentu Word"""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)
    
    serializer = DocumentUploadSerializer(data=request.data)
    if serializer.is_valid():
        uploaded = serializer.validated_data['file']
        name = serializer.validated_data.get('name', uploaded.name)

        # Obsługujemy tylko .docx (mammoth nie konwertuje .doc)
        if not uploaded.name.lower().endswith('.docx'):
            return Response({'error': 'Obsługiwane są tylko pliki .docx (zapisz dokument jako DOCX).'},
                            status=status.HTTP_400_BAD_REQUEST)

        # Konwertuj dokument Word do HTML
        try:
            data = uploaded.read()
            if not data:
                return Response({'error': 'Przesłany plik jest pusty.'}, status=status.HTTP_400_BAD_REQUEST)
            result = mammoth.convert_to_html(io.BytesIO(data))
            html_content = result.value or ''
        except Exception as e:
            return Response({'error': f'Błąd przetwarzania pliku DOCX: {str(e)}'}, 
                            status=status.HTTP_400_BAD_REQUEST)

        # Stwórz dokument w bazie
        content_file = ContentFile(data)
        content_file.name = uploaded.name
        document = Document.objects.create(
            name=name,
            file=content_file,
            original_content=html_content,
            created_by=request.user
        )

        serializer = DocumentSerializer(document)
        return Response(serializer.data, status=status.HTTP_201_CREATED)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def reprocess_document(request, document_id: int):
    """Ponownie przetwórz istniejący dokument DOCX do HTML (tylko admin, tylko własne dokumenty)."""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)

    try:
        document = Document.objects.get(id=document_id, created_by=request.user)
    except Document.DoesNotExist:
        return Response({'error': 'Dokument nie istnieje'}, status=status.HTTP_404_NOT_FOUND)

    file = document.file
    if not file or not file.name.lower().endswith('.docx'):
        return Response({'error': 'Obsługiwane są tylko pliki .docx (zapisz dokument jako DOCX).'},
                        status=status.HTTP_400_BAD_REQUEST)

    try:
        with file.open('rb') as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value or ''
        try:
            file.seek(0)
        except Exception:
            pass
        document.original_content = html_content
        document.save(update_fields=['original_content', 'updated_at'])
        serializer = DocumentSerializer(document)
        return Response(serializer.data)
    except Exception as e:
        return Response({'error': f'Błąd przetwarzania pliku DOCX: {str(e)}'},
                        status=status.HTTP_400_BAD_REQUEST)
    
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_field(request):
    """Stworzenie pola do edycji w dokumencie"""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)
    
    serializer = FieldCreationSerializer(data=request.data)
    if serializer.is_valid():
        data = serializer.validated_data
        
        try:
            document = Document.objects.get(id=data['document_id'], created_by=request.user)
        except Document.DoesNotExist:
            return Response({'error': 'Dokument nie istnieje'}, status=status.HTTP_404_NOT_FOUND)
        
        field = EditableField.objects.create(
            document=document,
            field_id=data['field_id'],
            label=data['label'],
            placeholder=data.get('placeholder', ''),
            field_type=data['field_type'],
            original_value=data.get('original_value', ''),
            position_start=data['position_start'],
            position_end=data['position_end']
        )
        
        serializer = EditableFieldSerializer(field)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_field(request, field_id: int):
    """Usuń istniejące pole do edycji (tylko admin i tylko dla własnego dokumentu)."""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)

    try:
        field = EditableField.objects.select_related('document').get(id=field_id)
    except EditableField.DoesNotExist:
        return Response({'error': 'Pole nie istnieje'}, status=status.HTTP_404_NOT_FOUND)

    if field.document.created_by_id != request.user.id:
        return Response({'error': 'Brak uprawnień do tego pola'}, status=status.HTTP_403_FORBIDDEN)

    field.delete()
    return Response({'success': True})


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def assign_document(request):
    """Przypisanie dokumentu do użytkowników"""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)
    
    serializer = AssignDocumentSerializer(data=request.data)
    if serializer.is_valid():
        data = serializer.validated_data
        
        try:
            document = Document.objects.get(id=data['document_id'], created_by=request.user)
        except Document.DoesNotExist:
            return Response({'error': 'Dokument nie istnieje'}, status=status.HTTP_404_NOT_FOUND)
        
        # Sprawdź czy dokument ma pola do edycji
        if not document.editable_fields.exists():
            return Response({'error': 'Dokument musi mieć przynajmniej jedno pole do edycji'}, 
                          status=status.HTTP_400_BAD_REQUEST)
        
        # Admin może przypisać do użytkowników z własnej sekcji lub do siebie samego
        admin_section = (user_profile.section or '').strip()
        if not admin_section:
            return Response({'error': 'Administrator nie ma ustawionej sekcji w profilu.'}, status=status.HTTP_400_BAD_REQUEST)

        assignments_created = []
        invalid_targets = []
        for user_id in data['user_ids']:
            try:
                user = User.objects.get(id=user_id)
            except User.DoesNotExist:
                invalid_targets.append(user_id)
                continue

            # Dozwolone: użytkownik z tej samej sekcji albo sam admin
            target_profile = getattr(user, 'userprofile', None)
            same_section = (target_profile and (target_profile.section or '').strip() == admin_section)
            is_self = user.id == request.user.id
            if not (same_section or is_self):
                invalid_targets.append(user_id)
                continue

            assignment, created = DocumentAssignment.objects.get_or_create(
                document=document,
                user=user,
                defaults={'status': 'pending'}
            )
            if created:
                assignments_created.append(assignment)
        
        # Zmień status dokumentu na wysłany
        document.status = 'sent'
        document.save()
        message = f'Dokument przypisano do {len(assignments_created)} użytkowników'
        if invalid_targets:
            message += f". Pominieto ID spoza sekcji: {','.join(map(str, invalid_targets))}"
        return Response({'success': True, 'message': message})
    
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def admin_documents(request):
    """Lista dokumentów administratora"""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)
    
    documents = Document.objects.filter(created_by=request.user).order_by('-created_at')
    serializer = DocumentSerializer(documents, many=True)
    return Response(serializer.data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def user_assignments(request):
    """Lista przypisań dla użytkownika"""
    assignments = DocumentAssignment.objects.filter(user=request.user).order_by('-assigned_at')
    serializer = DocumentAssignmentSerializer(assignments, many=True)
    return Response(serializer.data)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def submit_field_values(request):
    """Zapisanie wartości pól przez użytkownika"""
    serializer = SubmitFieldValuesSerializer(data=request.data)
    if serializer.is_valid():
        data = serializer.validated_data
        
        try:
            assignment = DocumentAssignment.objects.get(
                id=data['assignment_id'], 
                user=request.user
            )
        except DocumentAssignment.DoesNotExist:
            return Response({'error': 'Przypisanie nie istnieje'}, status=status.HTTP_404_NOT_FOUND)
        
        # Zapisz wartości pól
        for field_id, value in data['field_values'].items():
            try:
                field = EditableField.objects.get(
                    document=assignment.document,
                    field_id=field_id
                )
                field_value, created = FieldValue.objects.update_or_create(
                    assignment=assignment,
                    field=field,
                    defaults={'value': value}
                )
            except EditableField.DoesNotExist:
                continue
        
        # Zaktualizuj status przypisania
        if assignment.status == 'pending':
            assignment.status = 'in_progress'
            assignment.started_at = timezone.now()
        
        assignment.save()
        
        return Response({'success': True, 'message': 'Wartości zostały zapisane'})
    
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def complete_assignment(request, assignment_id):
    """Finalizacja wypełniania dokumentu"""
    try:
        assignment = DocumentAssignment.objects.get(
            id=assignment_id,
            user=request.user
        )
    except DocumentAssignment.DoesNotExist:
        return Response({'error': 'Przypisanie nie istnieje'}, status=status.HTTP_404_NOT_FOUND)
    
    # Sprawdź czy wszystkie pola zostały wypełnione
    required_fields = assignment.document.editable_fields.all()
    filled_fields = assignment.field_values.all()
    
    if required_fields.count() != filled_fields.count():
        return Response({'error': 'Nie wszystkie pola zostały wypełnione'}, 
                      status=status.HTTP_400_BAD_REQUEST)
    
    assignment.status = 'completed'
    assignment.completed_at = timezone.now()
    assignment.save()

    # Wygeneruj plik DOCX z wstawionymi wartościami pól i zapisz wersję
    try:
        _generate_assignment_docx_version(assignment)
    except Exception as e:
        # Nie blokuj kończenia, jeśli generowanie pliku się nie powiedzie
        print(f"[WARN] DOCX generation failed for assignment {assignment.id}: {e}")
    
    return Response({'success': True, 'message': 'Dokument został wysłany pomyślnie'})


def _replace_in_runs(runs, old: str, new: str):
    """Replace text possibly spanning multiple runs while preserving styles
    outside the replaced span.

    - Only the text inside the matched placeholder (old) will be unified (goes into
      the first affected run); text before and after keeps their original runs/styles.
    - If no match, returns False.
    """
    try:
        if not runs or not old:
            return False

        texts = [r.text or '' for r in runs]
        combined = ''.join(texts)
        start = combined.find(old)
        if start < 0:
            return False
        end = start + len(old)

        # Compute offsets per run: [start_pos, end_pos)
        offsets = []
        pos = 0
        for t in texts:
            offsets.append((pos, pos + len(t)))
            pos += len(t)

        # Find runs overlapping [start, end)
        in_span = [i for i, (s, e) in enumerate(offsets) if not (e <= start or s >= end)]
        if not in_span:
            return False

        first_i, last_i = in_span[0], in_span[-1]
        first_s, first_e = offsets[first_i]
        last_s, last_e = offsets[last_i]

        # Compute prefix/suffix relative to first/last affected runs
        first_run_text = texts[first_i]
        last_run_text = texts[last_i]

        prefix_len = max(0, start - first_s)
        suffix_start = max(0, end - last_s)

        if first_i == last_i:
            # Replacement fully inside a single run
            runs[first_i].text = (
                first_run_text[:prefix_len] + new + first_run_text[suffix_start:]
            )
        else:
            # Set first run: prefix + new
            runs[first_i].text = first_run_text[:prefix_len] + new
            # Clear middle runs
            for i in in_span[1:-1]:
                runs[i].text = ''
            # Set last run: suffix
            runs[last_i].text = last_run_text[suffix_start:]
        return True
    except Exception:
        return False


def _generate_assignment_docx_version(assignment: DocumentAssignment) -> DocumentVersion:
    """Utwórz plik DOCX z wartościami pól. Zachowujemy style, podmieniając tekst w istniejących runach.
    Uwaga: jeśli placeholder (original_value) został rozbity na wiele runów, prosta podmiana może go nie znaleźć.
    """
    doc_model = assignment.document
    file_field = doc_model.file
    if not file_field or not file_field.name.lower().endswith('.docx'):
        raise ValueError('Brak pliku DOCX do przetworzenia')

    # Mapuj field_id -> value i field.original_value
    values_by_field_id = {fv.field.field_id: fv.value for fv in assignment.field_values.select_related('field').all()}
    fields = list(doc_model.editable_fields.all())

    with file_field.open('rb') as f:
        doc = DocxDocument(f)

    # Podmień w paragrafach
    for para in doc.paragraphs:
        for field in fields:
            old = field.original_value or ''
            if not old:
                continue
            new = values_by_field_id.get(field.field_id)
            if new is None:
                continue
            _replace_in_runs(para.runs, old, new)

    # Podmień w tabelach
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for field in fields:
                        old = field.original_value or ''
                        if not old:
                            continue
                        new = values_by_field_id.get(field.field_id)
                        if new is None:
                            continue
                        _replace_in_runs(p.runs, old, new)

    # Zapisz do pamięci i do FileField
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    def _sanitize(name: str) -> str:
        name = name or ''
        return ''.join(ch if ch.isalnum() or ch in (' ', '-', '_') else '_' for ch in name).strip().replace(' ', '_')

    filename_base = doc_model.name.rsplit('.', 1)[0]
    safe_doc = _sanitize(filename_base)
    safe_user = _sanitize(assignment.user.username)
    # Prefer pattern: username__document.docx for easy sorting by user
    out_name = f"{safe_user}__{safe_doc}.docx"

    content = ContentFile(buf.read())
    version = DocumentVersion.objects.create(
        assignment=assignment,
        content='',
    )
    version.generated_file.save(out_name, content, save=True)
    return version


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_assignment_docx(request, assignment_id: int):
    """Zwróć wygenerowany plik DOCX dla przypisania. Jeśli nie istnieje, wygeneruj go."""
    try:
        assignment = DocumentAssignment.objects.get(id=assignment_id)
    except DocumentAssignment.DoesNotExist:
        return Response({'error': 'Przypisanie nie istnieje'}, status=status.HTTP_404_NOT_FOUND)

    # Dostęp: admin, który utworzył dokument, lub użytkownik będący właścicielem assignment
    is_admin_owner = assignment.document.created_by_id == request.user.id
    is_user_self = assignment.user_id == request.user.id
    if not (is_admin_owner or is_user_self):
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)

    version = assignment.versions.order_by('-created_at').first()
    if not version or not version.generated_file:
        try:
            version = _generate_assignment_docx_version(assignment)
        except Exception as e:
            return Response({'error': f'Błąd generowania pliku: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

    if not version.generated_file:
        return Response({'error': 'Brak wygenerowanego pliku'}, status=status.HTTP_404_NOT_FOUND)

    file = version.generated_file
    response = FileResponse(file.open('rb'), as_attachment=True, filename=file.name.split('/')[-1])
    return response


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_document(request, document_id: int):
    """Usuń dokument (tylko admin i tylko własne). Czyści również pliki wygenerowane oraz oryginalny plik."""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)

    try:
        document = Document.objects.get(id=document_id, created_by=request.user)
    except Document.DoesNotExist:
        return Response({'error': 'Dokument nie istnieje'}, status=status.HTTP_404_NOT_FOUND)

    with transaction.atomic():
        # Usuń wygenerowane pliki wersji powiązanych przypisań
        versions = DocumentVersion.objects.filter(assignment__document=document)
        for v in versions:
            if v.generated_file:
                try:
                    v.generated_file.delete(save=False)
                except Exception:
                    pass
        # Usuń plik źródłowy dokumentu
        if document.file:
            try:
                document.file.delete(save=False)
            except Exception:
                pass
        document.delete()

    return Response({'success': True})


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_assignment(request, assignment_id: int):
    """Usuń przypisanie (np. ukończony dokument). Admin może usuwać tylko przypisania swoich dokumentów."""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)

    try:
        assignment = DocumentAssignment.objects.select_related('document').get(id=assignment_id)
    except DocumentAssignment.DoesNotExist:
        return Response({'error': 'Przypisanie nie istnieje'}, status=status.HTTP_404_NOT_FOUND)

    if assignment.document.created_by_id != request.user.id:
        return Response({'error': 'Brak uprawnień do tego przypisania'}, status=status.HTTP_403_FORBIDDEN)

    with transaction.atomic():
        # Usuń wygenerowane pliki z wersji
        for v in assignment.versions.all():
            if v.generated_file:
                try:
                    v.generated_file.delete(save=False)
                except Exception:
                    pass
        assignment.delete()

    return Response({'success': True})


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def completed_assignments(request):
    """Lista ukończonych przypisań (dla adminów)"""
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)
    
    assignments = DocumentAssignment.objects.filter(
        document__created_by=request.user,
        status='completed'
    ).order_by('-completed_at')
    
    serializer = DocumentAssignmentSerializer(assignments, many=True)
    return Response(serializer.data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_completed_zip(request):
    """Zbiorcze pobranie ukończonych przypisań jako ZIP.
    Opcjonalnie przyjmujemy query param ?document_id=ID, aby ograniczyć do jednego dokumentu.
    Dostępne wyłącznie dla admina i tylko dla jego dokumentów.
    """
    user_profile = getattr(request.user, 'userprofile', None)
    if not user_profile or user_profile.role != 'admin':
        return Response({'error': 'Brak uprawnień'}, status=status.HTTP_403_FORBIDDEN)

    document_id = request.GET.get('document_id')
    assignments_qs = DocumentAssignment.objects.filter(
        document__created_by=request.user,
        status='completed'
    )
    if document_id:
        try:
            doc = Document.objects.get(id=int(document_id), created_by=request.user)
        except (Document.DoesNotExist, ValueError):
            return Response({'error': 'Dokument nie istnieje lub nie masz do niego dostępu'}, status=status.HTTP_404_NOT_FOUND)
        assignments_qs = assignments_qs.filter(document=doc)

    assignments = list(assignments_qs.order_by('document__name', 'user__username'))
    if not assignments:
        return Response({'error': 'Brak ukończonych przypisań do pobrania'}, status=status.HTTP_404_NOT_FOUND)

    # Zbuduj ZIP w pamięci
    buf = io.BytesIO()
    def _sanitize(name: str) -> str:
        name = name or ''
        return ''.join(ch if ch.isalnum() or ch in (' ', '-', '_') else '_' for ch in name).strip().replace(' ', '_')

    with zipfile.ZipFile(buf, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
        for ass in assignments:
            # Upewnij się, że wersja istnieje (wygeneruj jeśli brak)
            version = ass.versions.order_by('-created_at').first()
            if not version or not version.generated_file:
                try:
                    version = _generate_assignment_docx_version(ass)
                except Exception:
                    continue
            if not version or not version.generated_file:
                continue

            try:
                f = version.generated_file
                with f.open('rb') as fh:
                    data = fh.read()
                # Zbuduj nazwę pliku w ZIP
                base = _sanitize(ass.document.name.rsplit('.', 1)[0])
                uname = _sanitize(ass.user.username)
                # Prefer same pattern as generator
                arcname = f"{uname}/{uname}__{base}.docx"
                zf.writestr(arcname, data)
            except Exception:
                continue

    buf.seek(0)
    stamp = datetime.now().strftime('%Y%m%d_%H%M')
    if document_id:
        zip_base = _sanitize(doc.name.rsplit('.', 1)[0])
        zip_name = f"{zip_base}_completed_{stamp}.zip"
    else:
        zip_name = f"completed_{_sanitize(request.user.username)}_{stamp}.zip"
    response = FileResponse(buf, as_attachment=True, filename=zip_name)
    return response
